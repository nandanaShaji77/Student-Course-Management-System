import os
import sys
import subprocess

# Ensure python-docx library is installed on the host
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing python-docx library to compile Microsoft Word report...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background shading color for table cells."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner padding margins for table cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, title, code_text):
    """Adds a beautiful, professional monospaced code block container in the document."""
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(f"Code Snippet: {title}")
    run_title.bold = True
    run_title.font.size = Pt(9.5)
    run_title.font.color.rgb = RGBColor(124, 58, 237) # Dark Violet Brand Color
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_background(cell, "F8FAFC") # Soft slate background
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    # Border formatting (thin solid grey left-accented line)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="18" w:space="0" w:color="7C3AED"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    # Fill source code
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59) # Slate color
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_screenshot_placeholder(doc, label):
    """Creates a beautifully styled placeholder frame for inserting laboratory screenshots."""
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_background(cell, "F1F5F9") # Muted gray background
    set_cell_margins(cell, top=250, bottom=250, left=200, right=200)
    
    # Thin dashed border for screenshot insertion
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="dashed" w:sz="6" w:color="94A3B8"/>
            <w:left w:val="dashed" w:sz="6" w:color="94A3B8"/>
            <w:bottom w:val="dashed" w:sz="6" w:color="94A3B8"/>
            <w:right w:val="dashed" w:sz="6" w:color="94A3B8"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    
    run_label = p.add_run(f"[ Insert Academic Screenshot Here: {label} ]\n")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(71, 85, 105)
    
    run_sub = p.add_run("(Right-click this box -> Add Image to replace this placeholder with your actual execution output)")
    run_sub.italic = True
    run_sub.font.size = Pt(8.5)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def load_short_file_content(filepath, max_lines=35, fallback_msg="// Code snippet not available."):
    """Reads a concise, readable subset of the workspace file to avoid report clutter."""
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                if len(lines) > max_lines:
                    return "".join(lines[:max_lines]) + "\n\n// ... [Code truncated for brevity; full file compiled in workspace] ..."
                return "".join(lines)
        except Exception as e:
            return f"// Error reading file {os.path.basename(filepath)}: {str(e)}"
    return f"// File missing: {filepath}\n{fallback_msg}"

def generate_report():
    doc = docx.Document()
    
    # Set standard academic 1-inch page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Standard formatting style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(15, 23, 42) # Slate black
    
    # ==========================================
    #             1. COVER PAGE
    # ==========================================
    
    # Document Title
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.paragraph_format.space_before = Pt(40)
    p_head.paragraph_format.space_after = Pt(20)
    run_head = p_head.add_run("PROJECT REPORT")
    run_head.bold = True
    run_head.font.size = Pt(24)
    run_head.font.name = 'Arial'
    run_head.font.color.rgb = RGBColor(15, 23, 42)
    
    for _ in range(3):
        doc.add_paragraph()
        
    # Project Title (shaded block matching mockup)
    tbl_name = doc.add_table(rows=1, cols=1)
    tbl_name.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_name = tbl_name.cell(0, 0)
    cell_name.width = Inches(5.5)
    set_cell_background(cell_name, "F1F5F9") # Slate 100 bg
    set_cell_margins(cell_name, top=200, bottom=200, left=250, right=250)
    
    tcPr = cell_name._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:color="CBD5E1"/>
            <w:left w:val="single" w:sz="6" w:color="CBD5E1"/>
            <w:bottom w:val="single" w:sz="6" w:color="CBD5E1"/>
            <w:right w:val="single" w:sz="6" w:color="CBD5E1"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p_name = cell_name.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("Student Course Management System")
    run_name.bold = True
    run_name.font.size = Pt(18)
    run_name.font.name = 'Arial'
    run_name.font.color.rgb = RGBColor(71, 85, 105) # Slate grey
    
    for _ in range(5):
        doc.add_paragraph()
        
    # University Logo
    logo_path = r"C:\Users\jitna\.gemini\antigravity\brain\114d0866-8751-474f-805f-480389cfee7a\media__1779429945863.png"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(logo_path, width=Inches(3.2))
    else:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p_logo.add_run("[ LOVELY PROFESSIONAL UNIVERSITY ]")
        run_l.bold = True
        run_l.font.size = Pt(14)
        run_l.font.color.rgb = RGBColor(234, 88, 12)

    doc.add_page_break()

    # ==========================================
    #             2. REPORT BODY
    # ==========================================
    
    report_sections = [
        {
            "title": "1. Introduction",
            "text": (
                "Managing academic catalogs, course records, and enrollment transactions represents a standard operational "
                "task in educational environments. In traditional setups, applications are run directly on host servers, "
                "which introduces environment-specific issues, port conflicts, and installation overhead.\n\n"
                "This project demonstrates an academic deployment of a Student Course Management System built on a "
                "containerized multi-container architecture. By separating the database tier, backend API application, "
                "and static web gateway into independent containers, the system achieves environment independence. "
                "Furthermore, the project integrates automated compilation, packaging, and CI/CD pipelines (via Maven, "
                "GitHub Actions, and Jenkins) to teach and demonstrate modern DevOps principles and deployment automation."
            )
        },
        {
            "title": "2. Problem Statement",
            "text": (
                "Standard educational web applications face several practical issues during local setup and deployment:\n"
                "• Host Environment Inconsistencies: Discrepancies in Java JDK runtimes, local database schemas, or package versions prevent code from building uniformly across developer machines.\n"
                "• Cross-Origin Resource Sharing (CORS): Direct communication between a frontend portal and an API backend across different ports triggers browser CORS blocks, complicating request handling.\n"
                "• Port Binding Conflicts: Applications binding directly to common host ports (like 80 or 8080) clash with existing developer software, causing runtime errors.\n"
                "• Manual Packaging Bottlenecks: Compiling code, running unit tests, and configuring dependencies manually slows down deployment and increases error rates."
            )
        },
        {
            "title": "3. Objective",
            "text": (
                "The core objective of this academic B.Tech project is to resolve host dependencies and automate the "
                "deployment lifecycle. The project aims to:\n"
                "1. Containerize the applications using Docker to separate the PostgreSQL database, Java Spring Boot API, and Nginx web proxy.\n"
                "2. Standardize local networking and database volumes using Docker Compose to prevent port binding issues.\n"
                "3. Configure Nginx to resolve CORS issues by proxying client REST requests to the internal API network.\n"
                "4. Automate building and unit test verification using Maven lifecycle stages.\n"
                "5. Create a GitHub Actions pipeline to compile source code, execute tests, and push container images on repository updates.\n"
                "6. Write a declarative Jenkinsfile CD pipeline to automate checkout, testing, building, and deployment."
            )
        },
        {
            "title": "4. Proposed System",
            "text": (
                "The proposed Student Course Management System introduces a 3-Tier Containerized Multi-Container Architecture:\n"
                "1. Presentation Tier (Nginx Proxy): An Alpine-based Nginx container that hosts the static HTML5/CSS/JS frontend dashboard. It maps incoming client calls on port 8082 and routes them internally to the API container.\n"
                "2. Logic Tier (Spring Boot): A Java 17 Spring Boot service inside a JRE container that handles database CRUD operations for student and course catalogs.\n"
                "3. Database Tier (PostgreSQL): A PostgreSQL 15 container storing records with persistent volume mapping.\n\n"
                "Relational Model: Implements a Many-to-Many relationship between Student and Course tables, joined via a junction table `student_courses` to prevent loops during JSON serialization."
            )
        },
        {
            "title": "5. System Architecture Diagram",
            "text": (
                "The diagram below describes the multi-container environment where Nginx acts as the single entry point "
                "for all client calls, routing them inside the containerized network:"
            ),
            "ascii_diagram": (
                " +---------------------------------------------------------+\n"
                " |                       Client Layer                      |\n"
                " |                     (Browser Client)                    |\n"
                " +----------------------------+----------------------------+\n"
                "                              | HTTP (Host Port 8082)\n"
                "                              v\n"
                " +---------------------------------------------------------+\n"
                " |                    Web/Proxy Gateway                    |\n"
                " |                (Nginx Proxy Container)                  |\n"
                " +---------+-------------------------------------+---------+\n"
                "           | Serves Static Files                 | Proxies /api/* requests\n"
                "           v                                     v\n"
                " +-------------------+                 +-------------------+\n"
                " | Frontend Assets   |                 | Spring Boot App   |\n"
                " | (HTML, CSS, JS)   |                 | (API Container)   |\n"
                " +-------------------+                 +---------+---------+\n"
                "                                                 | JDBC (Port 5432)\n"
                "                                                 v\n"
                "                                        +-------------------+\n"
                "                                        | PostgreSQL DB     |\n"
                "                                        | (DB Container)    |\n"
                "                                        +---------+---------+\n"
                "                                                  | Persists Data\n"
                "                                                  v\n"
                "                                        +-------------------+\n"
                "                                        | Persistent Volume |\n"
                "                                        | (postgres_data)   |\n"
                "                                        +-------------------+\n"
            )
        },
        {
            "title": "6. Methodology",
            "text": (
                "The academic implementation splits development into standard structured phases:\n"
                "• Local Development: Coding the Java Spring Boot models, repositories, and web controllers alongside the frontend.\n"
                "• Build Automation: Utilizing the Maven compiler and surefire plug-in to compile code and run tests offline without database dependecies.\n"
                "• Containerization: Writing Dockerfiles for Nginx and Spring Boot to build lightweight, multi-stage images.\n"
                "• Multi-container Orchestration: Writing a `docker-compose.yml` to set up service startups, network links, and volume persistence.\n"
                "• Automated CI Pipeline: Setting up GitHub Actions to run builds and test suites, and push Docker images to registries.\n"
                "• Automated CD Pipeline: Implementing a Jenkinsfile pipeline to automate compilation, testing, and compose deployments."
            )
        },
        {
            "title": "7. CI/CD Workflow Diagram",
            "text": (
                "The automation flow handles codebase commits, compiles and tests code, builds images, "
                "and automates local compose deployments as shown below:"
            ),
            "ascii_diagram": (
                " +---------------------------------------------------------+\n"
                " |                      Local Commit                       |\n"
                " |              (Git Push to main branch)                  |\n"
                " +----------------------------+----------------------------+\n"
                "                              | Triggers CI Pipeline\n"
                "                              v\n"
                " +---------------------------------------------------------+\n"
                " |                 Continuous Integration                  |\n"
                " |                 (GitHub Actions CI)                     |\n"
                " |  - Checks out codebase                                  |\n"
                " |  - Configures Java 17 Runtime                           |\n"
                " |  - Runs Compilation & JUnit Tests via Maven             |\n"
                " |  - Builds multi-stage Docker container images           |\n"
                " |  - Publishes production images to Docker Hub registry    |\n"
                " +----------------------------+----------------------------+\n"
                "                              | Triggers CD Pipeline\n"
                "                              v\n"
                " +---------------------------------------------------------+\n"
                " |                  Continuous Delivery                    |\n"
                " |                  (Jenkins CD Pipeline)                  |\n"
                " |  - Clones source code automatically                     |\n"
                " |  - Triggers offline Maven packaging & test cycles       |\n"
                " |  - Rebuilds container images locally                    |\n"
                " |  - Runs automated Docker Compose deployment updates     |\n"
                " +---------------------------------------------------------+\n"
            )
        },
        {
            "title": "8. Tools & Technologies",
            "text": (
                "The project utilizes core industry-standard DevOps tools aligned with the academic syllabus:\n"
                "• Spring Boot: Acts as the backend API container logic engine.\n"
                "• Maven: Automates compilation and dependencies using target pom settings.\n"
                "• Docker: Containerizes the database, API backend, and frontend portal.\n"
                "• Docker Compose: Coordinates multi-container networks, environments, and data volumes.\n"
                "• Nginx: Serves the static client assets and handles reverse-proxy routing.\n"
                "• GitHub Actions: Implements the Continuous Integration pipeline.\n"
                "• Jenkins: Automates the Continuous Delivery deployment pipeline."
            )
        },
        {
            "title": "9. Expected Outcome",
            "text": (
                "Upon running the multi-container configuration, the system produces the following results:\n"
                "• Automated Deployment: Developers can run `docker compose up --build -d` to launch the whole stack without manual server installation.\n"
                "• Port Binding Safety: By mapping the services to open ports (3000 mapped to 8082, and 8080 to 8081), host port conflicts are prevented.\n"
                "• Stable Persistence: PostgreSQL data persists in the mapped volume even if containers are stopped or deleted.\n"
                "• Responsive Dashboard: Students, courses, and registrations are managed dynamically on an interactive panel."
            )
        },
        {
            "title": "10. Limitations",
            "text": (
                "While the system fulfills academic DevOps objectives, it has a few practical limitations:\n"
                "• Local Orchestration Constraints: Running multi-containers locally depends on Docker Desktop, which requires significant host memory and CPU resources.\n"
                "• Basic Database Authentication: The PostgreSQL container uses basic environment credentials without external vault integrations.\n"
                "• Single-Host Execution: Docker Compose is restricted to a single host server and lacks multi-node scaling and replication features found in Kubernetes."
            )
        },
        {
            "title": "11. Future Enhancements",
            "text": (
                "To extend the project, several enhancements can be explored:\n"
                "• Kubernetes Migration: Transitioning the Docker Compose spec into Kubernetes (K8s) deployment manifests to enable scaling.\n"
                "• Automated Testing in Pipelines: Integrating integration testing tools (like Selenium or Playwright) into the pipeline to run UI tests.\n"
                "• Secured Secrets Handling: Utilizing HashiCorp Vault or environment secret keys to keep PostgreSQL credentials out of source code."
            )
        },
        {
            "title": "12. Conclusion",
            "text": (
                "The Student Course Management System DevOps project provides a hands-on implementation of automated deployment pipelines. "
                "By containerizing the application layers, the system avoids host dependencies and port binding issues. "
                "Automating tasks with Maven, GitHub Actions, and Jenkins illustrates how modern software development pipelines "
                "ensure code quality, run automated tests, and streamline delivery."
            )
        }
    ]

    for sec in report_sections:
        # Heading Style
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(14)
        p_head.paragraph_format.space_after = Pt(4)
        p_head.paragraph_format.keep_with_next = True
        run_h = p_head.add_run(sec["title"])
        run_h.bold = True
        run_h.font.size = Pt(13)
        run_h.font.color.rgb = RGBColor(15, 23, 42)
        
        # Paragraph Text
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_after = Pt(6)
        p_body.paragraph_format.line_spacing = 1.15
        run_b = p_body.add_run(sec["text"])
        run_b.font.size = Pt(10.5)
        
        # Add ASCII Diagram if available
        if "ascii_diagram" in sec:
            add_code_block(doc, f"{sec['title']} (Visual Diagram)", sec["ascii_diagram"])
            
        # Add filtered, concise code snippets matching B.Tech topics
        if "Proposed System" in sec["title"]:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            student_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\backend\src\main\java\com\example\studentcourse\model\Student.java", max_lines=28)
            add_code_block(doc, "Student.java (JPA Model Relational Annotations)", student_code)
            
            course_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\backend\src\main\java\com\example\studentcourse\model\Course.java", max_lines=28)
            add_code_block(doc, "Course.java (JPA Model Relational Annotations)", course_code)
            
        elif "Methodology" in sec["title"]:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            controller_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\backend\src\main\java\com\example\studentcourse\controller\StudentController.java", max_lines=30)
            add_code_block(doc, "StudentController.java (REST API Mappings)", controller_code)
            
            nginx_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\frontend\nginx.conf", max_lines=25)
            add_code_block(doc, "nginx.conf (CORS Mitigation Reverse Proxy)", nginx_code)
            
            app_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\frontend\app.js", max_lines=25)
            add_code_block(doc, "app.js (Javascript AJAX and Metric Calculations)", app_code)

        elif "Tools & Technologies" in sec["title"]:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            compose_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\docker-compose.yml", max_lines=35)
            add_code_block(doc, "docker-compose.yml (Multi-Container Orchestration)", compose_code)

        elif "Expected Outcome" in sec["title"]:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            ci_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\.github\workflows\ci.yml", max_lines=30)
            add_code_block(doc, "ci.yml (GitHub Actions Compilation & Tests)", ci_code)
            
            jenkins_code = load_short_file_content(r"c:\Users\jitna\OM2 Main Project\Jenkinsfile", max_lines=35)
            add_code_block(doc, "Jenkinsfile (Linux sh Continuous Delivery Pipeline)", jenkins_code)

    # ==========================================
    #             3. SCREENSHOTS SECTION
    # ==========================================
    doc.add_page_break()
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_before = Pt(14)
    p_head.paragraph_format.space_after = Pt(8)
    run_h = p_head.add_run("13. Execution Screenshots Section")
    run_h.bold = True
    run_h.font.size = Pt(14)
    run_h.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run(
        "This section contains academic execution screenshots confirming the successful containerization, "
        "compilation, pipeline validation, and deployment operations of the B.Tech project:"
    )
    run_sub.font.size = Pt(10.5)
    
    screenshots = [
        "Frontend Dashboard - Dynamic Academic Analytics Panel",
        "Docker Containers Running - Active PostgreSQL, Backend, Nginx Containers",
        "GitHub Actions Success - Automated Compiles and Image Push Pipelines",
        "Jenkins Pipeline Success - Automated Pipeline Continuous Delivery Stages",
        "Docker Compose Deployment - Successful Networking Bindings on Ports 8082/8081"
    ]
    
    for scr in screenshots:
        add_screenshot_placeholder(doc, scr)

    # Save finalized report
    output_path = r"c:\Users\jitna\OM2 Main Project\Student_Course_Management_System_Project_Report.docx"
    doc.save(output_path)
    print(f"Academic Report compiled successfully at: {output_path}")

if __name__ == "__main__":
    generate_report()
